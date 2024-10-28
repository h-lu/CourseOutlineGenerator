import streamlit as st
import json
from openai import OpenAI
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE  # 添加这行导入
import io
import datetime

# 设置页面配置必须是第一个 Streamlit 命令
st.set_page_config(page_title="课程考试生成器", page_icon="📚", layout="wide")

# 在文件开头添加这个函数定义
def get_project_score_standards(project_requirements):
    """生成项目评分标准文本"""
    standards = {
        "需求分析": {
            "score": 20,
            "items": [
                {"name": "需求完整性", "score": 5, "criteria": "需求覆盖度和准确性"},
                {"name": "分析深度", "score": 5, "criteria": "问题分析的深度和合理性"},
                {"name": "文档质量", "score": 5, "criteria": "文档的规范性和完整性"},
                {"name": "创新性", "score": 5, "criteria": "解决方案的创新程度"}
            ]
        },
        "概要设计": {
            "score": 25,
            "items": [
                {"name": "架构设计", "score": 8, "criteria": "系构的合理性和可扩展性"},
                {"name": "技术方案", "score": 7, "criteria": "技术选型的适当性"},
                {"name": "模块划分", "score": 5, "criteria": "模块划分的清晰度和耦合度"},
                {"name": "创新性", "score": 5, "criteria": "设计方案的创新性"}
            ]
        },
        "详细设计": {
            "score": 20,
            "items": [
                {"name": "设计完整性", "score": 6, "criteria": "设计文档的完整性"},
                {"name": "设计合理性", "score": 6, "criteria": "设计方案的可行性"},
                {"name": "规范性", "score": 4, "criteria": "设计规范的遵循程度"},
                {"name": "创新性", "score": 4, "criteria": "设计细节的创新点"}
            ]
        },
        "代码实现": {
            "score": 25,
            "items": [
                {"name": "功能实现", "score": 8, "criteria": "功能的完整性和正确性"},
                {"name": "代码质量", "score": 7, "criteria": "代码规范性和可维护性"},
                {"name": "性能优化", "score": 5, "criteria": "代码执行效率"},
                {"name": "新实现", "score": 5, "criteria": "技术实现的创新性"}
            ]
        },
        "测试报告": {
            "score": 10,
            "items": [
                {"name": "测试覆盖", "score": 3, "criteria": "测试用例的覆盖度"},
                {"name": "测试执行", "score": 3, "criteria": "测试执行的完整性"},
                {"name": "问题修复", "score": 2, "criteria": "问题跟踪和解决情况"},
                {"name": "报告质量", "score": 2, "criteria": "测试报告的规范性"}
            ]
        }
    }
    
    result = []
    total_score = 0
    
    for req in project_requirements:
        if req in standards:
            std = standards[req]
            total_score += std['score']
            result.append(f"\n{req}（{std['score']}分）：")
            for item in std['items']:
                result.append(f"- {item['name']}（{item['score']}分）：{item['criteria']}")
    
    # 添加总分说明
    if result:
        result.insert(0, f"总分：{total_score}分")
    
    return "\n".join(result)

def load_course_outline(uploaded_file):
    """加载课程大纲JSON文件"""
    if uploaded_file is not None:
        return json.load(uploaded_file)
    return None

def create_exam_prompt(outline_data, exam_type, chapters=None, additional_requirements=None, config=None):
    """创建考试提示"""
    # 获取课程基本信息
    course_type = outline_data['basic_info']['course_type']
    department = outline_data['basic_info']['department']
    major = outline_data['basic_info']['major']
    
    # 处理章节文本
    if chapters and isinstance(chapters, list):
        chapters_text = ', '.join(chapters)
    else:
        chapters_text = "全部章节"
    
    # 根据考试类型选择JSON格式
    if exam_type == "实验":
        json_format = get_experiment_json_format()
    elif exam_type == "大作业":
        json_format = get_project_json_format()
    else:
        json_format = '''{
            "questions": [
                {
                    "type": "选择题/判断题/简答题/编程题",
                    "question": "题目内容",
                    "options": ["A. 选项1", "B. 选项2", "C. 选项3", "D. 选项4"],
                    "answer": "标准答案",
                    "explanation": "解题思路和解析",
                    "course_objectives": ["对应的课程目标"],
                    "aacsb_goals": ["对应的AACSB目标"],
                    "difficulty": "基础/中等/困",
                    "score": "分值"
                }
            ]
        }'''
    
    system_prompt = f"""
    你一位专业的教育考试出题专家，门负责{department}{major}专业的{course_type}考试命题工作
    请于提供的课程信息生成高质量的考试题目，并以JSON格式输出，要求如下：

    1. 考试目标：
       - 准确评学生对课程知识点的掌握程度
       - 考核学生的实践应用能力
       - 符合AACSB认证标准的要求
       - 体现专业特色和课程定位
    
    2. 题目要求：
       - 每道题目必须对应具体的课程目标和AACSB目标
       - 题目难度要合理分布，并明确标注难度级别
       - 实践类题目必须符合专业特点
       - 所有题都要提供详细的解析和评分标准
       - 确保题目符合{department}{major}专业的特点和要求
    
    3. 知识点关联要求：
       - 体现知识点间的逻辑关系和递进性
       - 注重前后章节知识的连贯性
       - 强调知识点的实际应用场景
       - 突出重点和难点知识的关联
       - 体现知识体系的完整性
    
    4. 能力递进要求：
       - 基础能力：概念理解和基本应用
       - 进阶能力：知识综合和问题分析
       - 高阶能力：创新设计和优化改进
       - 实践能力：工程实践和项目开发
    
    5. JSON输出格式：
    {json_format}
    
    请确保输出严格遵循上述JSON格式。
    """
    
    # 根据不同考试类型调整提示内容
    exam_type_prompts = {
        "练习": f"""
        {course_type}练习题要求：
        
        练习难度要求（{config.get('difficulty', '中等')}）：
        {
            '''
            基础难度（Remember & Understand）：
            1. 题型分布：
               - 基础概念题（40%）：考察核心概念理解
               - 简单应用题（40%）：考察基本原理应用
               - 基础分析题（20%）：考察简单问题分析
            
            2. 知识覆盖：
               - 核心概念准确理解
               - 基本原理正确应用
               - 标准流程熟练掌握
               - 基础编程技能运用
            
            3. 答题要求：
               - 概念定义准确清晰
               - 基本步骤完整规范
               - 代码实现基本正确
               - 结果计算准确无误
            
            4. 评分标准：
               - 概念理解（40%）：术语准确、原理清晰
               - 应用能力（40%）：步骤正确、结果准确
               - 表达���力（20%）：条理清晰、书写规范
            ''' if config.get('difficulty') == '基础' else
            '''
            中等难度（Apply & Analyze）：
            1. 题型分布：
               - 概念应用题（30%）：考察知识综合运用
               - 综合分析题（40%）：考察问题分析能力
               - 实践应用题（30%）：考察实际问题解决
            
            2. 知识覆盖：
               - 多知识点综合运用
               - 实际问题分析能力
               - 算法设计与优化
               - 代码实现与调试
            
            3. 答题要求：
               - 分析过程逻辑完整
               - 解决方案合理可行
               - 代码实现规范高效
               - 结果分析深入准确
            
            4. 评分标准：
               - 分析能力（35%）：思路清晰、逻辑严谨
               - 解决方案（35%）：方案合理、实现正确
               - 创新思维（30%）：方法创新、优化改进
            ''' if config.get('difficulty') == '中等' else
            '''
            提高难度（Evaluate & Create）：
            1. 题型分布：
               - 综合应用题（30%）：考察系统性思维
               - 方案设计题（40%）：考察创新设计能力
               - 优化改进题（30%）：考察问题优化能力
            
            2. 知识覆盖：
               - 系统架构设计能力
               - 性能优化分析能力
               - 创新解决方案能力
               - 工程实践应用能力
            
            3. 答题要求：
               - 方案设计完整创新
               - 性能分析深入全面
               - 代码实现高效优雅
               - 优化思路清晰可行
            
            4. 评分标准：
               - 设计能力（30%）：方案创新、架构合理
               - 实现能力（30%）：代码优秀、性能优化
               - 创新能力（40%）：思路创新、解决方案优秀
            '''
        }
        
        通用评分标准：
        1. 答题完整性（20%）：
           - 问题理解准确
           - 解答步骤完整
           - 结果表达清晰
        
        2. 专业规范性（30%）：
           - 专业术语使用准确
           - 解题思路专业规范
           - 代码编写符合规范
        
        3. 创新思维（20%）：
           - 解题思路创新
           - 方法选择合理
           - 优化改进意识
        
        4. 实践应用（30%）：
           - 理论联系实际
           - 解决实际问题
           - 考虑实际约束
        
        要求说明：
        1. 题型要求（必须严格按照以下数量生成）：
        {chr(10).join([
            f"   - {q_type}：必须生成{details['count']}题（{details['description']}）"
            for q_type, details in config.get('practice_types', {}).items()
        ])}
        
        2. 知识点覆盖：
           - 重点章节内容占比70%
           - 基础知识考察30%
           - 确保知识点分布合理
        
        3. 题目要求：
           - 每道题目必须标注对应的课程目标和AACSB目标
           - 每道题目必须包含详细的答案和解析
           - 题目描述必须清晰准确
           - 答案和评分标准必须明确
        
        4. 实施要求：
           - 题目数量必须严格按照要求生成
           - 每道题的分值必须与设定一致
           - 各种题型的总分必须符合设定
           - 必须覆盖选定的所有章节内容
           - 确保题目难度分布合理
        
        5. 特殊说明：
           - 选择题必须包含4个选项
           - 判断题必须明确是非
           - 简答题必须提供详细评分点
           - 编程题必须包含测试用例
        """,
        
        "实验": f"""
        请生成一个完整的{course_type}实验指导方案，类型为{config.get('lab_type', '综合性')}实验。
        
        实验难度要求（{config.get('difficulty', '中等')}）：
        {
            '''
            基础难度（Remember & Understand）：
            1. 实验内容：
               - 聚焦单个核心知识点
               - 步骤清晰，流程明确
               - 提供完整的代码框架
               - 详细的操作指导
            
            2. 实验要求：
               - 基本功能的实现
               - 代码规范性要求
               - 简单的功能测试
               - 基础实验报告
            
            3. 评分标准：
               - 功能完整性（40%）
               - 代码规范性（30%）
               - 实验报告（30%）
            ''' if config.get('difficulty') == '基础' else
            '''
            中等难度（Apply & Analyze）：
            1. 实验内容：
               - 组合多个知识点
               - 部分步骤需自主设计
               - 提供部分代码框架
               - 关键步骤有指导
            
            2. 实验要求：
               - 完整功能实现
               - 代码优化要求
               - 完整的测试方案
               - 详细的分析报告
            
            3. 评分标准：
               - 功能实现（35%）
               - 代码质量（35%）
               - 实验报告（30%）
            ''' if config.get('difficulty') == '中等' else
            '''
            提高难度（Evaluate & Create）：
            1. 实验内容：
               - 综合性问题解决
               - 完全自主设计
               - 仅提供基本框架
               - 鼓励创新设计
            
            2. 实验要求：
               - 创新性解决方案
               - 性能优化要求
               - 完整的测试体系
               - 深入的分析报告
            
            3. 评分标准：
               - 方案创新性（30%）
               - 实现质量（40%）
               - 分析报告（30%）
            '''
        }
        
        实验设计要求：
        1. 实验定位：
           {config.get('lab_type', '综合性')}实验的特点：
           {
               '''
               基础性实验：
               - 针对单个知识点的掌握和应用
               - 提供详细的操作指导
               - 给出完整的代码框架
               - 设置明确的检查点
               - 预期结果清晰具体
               ''' if config.get('lab_type') == '基础性' else
               '''
               综合性实验：
               - 涉及多个知识点的综合应用
               - 提供关键步骤的指导
               - 给出部分代码框架
               - 需要自主设计部分内容
               - 有一定的探索空间
               ''' if config.get('lab_type') == '综合性' else
               '''
               设计性实验：
               - 提供开放性问题
               - 需要自主设计解决方案
               - 只提基本框或不提供
               - 鼓励创新和多样化
               - 重视方案的可行性
               '''
           }

        2. 识要求：
           - 应章节：{chapters_text}
           - 涉及知识点：请根章节内容列
           - 前置知识要求：请明确指出
        
        3. 实验内容：
           - 实验名称：应明确反映实验内容和类型
           - 实验目标：应对应课程目标和AACSB目标
           - 实验步骤：应符合实验型特点
           - 预期结果：应明确可验证
        
        4. 实验指导：
           - 环境准备：详细的环境配置说明
           - 操作步骤：符合实验类型的详细程度
           - 代码模板：根据实验类型提供相应的代码
           - 注意事项：可能遇到的问题和解决方案
        
        5. 考核要求：
           - 实验准备（20分）：
             * 环境配置完整性
             * 预习报告质量
           - 实验实现（50分）：
             * 功能完成度
             * 代码质量
             * 实现效果
           - 实验报告（30分）：
             * 文档规范性
             * 结果分析
             * 总结反思
        
        6. 创新与拓展：
           - 提供选做内容
           - 鼓励创新维
           - 指出扩展方向
        
        请确保生成的实验内容：
        1. 符合{major}专业特点
        2. 难度适合{course_type}水平
        3. 符合{config.get('lab_type', '综合性')}实的特点
        4. 与所选章节内容紧密相关
        5. 实验步骤清晰可执行
        6. 分标准客观可量化
        
        请严格按照提供的JSON格式生成实验内容。
        """,
        
        "大作业": f"""
        {course_type}大作业要求：
        
        项目难度要求（{config.get('difficulty', '中等')}）：
        {
            '''
            基础难度（Remember & Understand）：
            1. 项目范围：
               - 单一功能模块开发
               - 清晰的需求定义
               - 基础技术栈应用
               - 标准开发流程
            
            2. 技术要求：
               - 基本功能实现
               - 标准代码规范
               - 基本错误处理
               - 简单的用户界面
            
            3. 文档要求：
               - 基础需求文档
               - 简单设计说明
               - 基本测试报告
               - 用户使用手册
            ''' if config.get('difficulty') == '基础' else
            '''
            中等难度（Apply & Analyze）：
            1. 项目范围：
               - 多功能模块集成
               - 完整的业务流程
               - 多样化技术应用
               - 规范化项目管理
            
            2. 技术要求：
               - 完整功能实现
               - 性能优化考虑
               - 安全性处理
               - 良好的用户体验
            
            3. 文档要求：
               - 详细设计文档
               - 完整测试方案
               - 性能分析报告
               - 项目总结报告
            ''' if config.get('difficulty') == '中等' else
            '''
            提高难度（Evaluate & Create）：
            1. 项目范围：
               - 完整系统开发
               - 创新性解决方案
               - 先进技术应用
               - 工程化项目管理
            
            2. 技术要求：
               - 创新功能实现
               - 高性能架构
               - 完整安全方案
               - 优秀用户体验
            
            3. 文档要求：
               - 架构设计文档
               - 完整技术方案
               - 系统性能报告
               - 项目创新说明
            '''
        }
        
        项目设计：
           - 结合{major}专业特点的任务
           - 符合行业实际需求
           - 技术路线建议
           - 开发规范指导
        
        项目要求：
           选定的项目要求：{', '.join(config.get('project_requirements', ['需求分析']))}
           
           具体要求：
           {chr(10).join([
               f"- {req}：" + (
                   '''
                   需求分析：
                   * 完整的需调研报告
                   * 用户需求析文档
                   * 功能需求规格说明
                   * 非功能性需求说明
                   * 系统用例图和描述
                   * 业务流程分析
                   ''' if req == "需求分析" else
                   '''
                   概要设计：
                   * 系统架构设计方案
                   * 技术选型说明
                   * 系统模块划分
                   * 数据库概要设计
                   * 接口设计规范
                   * 系统部署架构
                   ''' if req == "概要设计" else
                   '''
                   细设计：
                   * 详细的类图设计
                   * 完整的ER图设计
                   * 数据库表结构设计
                   * API接口详细说明
                   * 核心算法设
                   * 安全方案设计
                   ''' if req == "详细设计" else
                   '''
                   代码实现：
                   * 遵循编码规范
                   * 完注释文档
                   * 代码版本控制
                   * 单元测试用例
                   * 代码审查记录
                   * 性能优化方
                   ''' if req == "代码实现" else
                   '''
                   测试报告：
                   * 测试计划文档
                   * 测试用例设计
                   * 测试执行记录
                   * 缺陷跟踪报告
                   * 性能测试报告
                   * 安全测试报告
                   '''
               ) for req in config.get('project_requirements', ['需求分析'])
           ])}
        
        评分标准（总分100分）：
           {get_project_score_standards(config.get('project_requirements', ['需求分析']))}
        
        提交要求：
           - 提交时间：严格遵守截止日期
           - 提交格式：
             * 源代码（包含完整的项目文件）
             * 文档材料（PDF格式）
             * 演示文稿（PPT格式）
             * 项目演示视频（可
         
        加分项（总分基础上最多加10分）：
           - 技术创新：使用新技术或创新解决方案（+3分）
           - 实用价值：具有际应价值（+3分）
           - 完整性文档完整、结构清晰（+2分）
           - 答辩表现：项目展示和答辩优秀（+2分）
        
        扣分项：
           - 迟交：每迟交一天扣总5分
           - 抄袭：发现抄袭直接记0分
           - 文档缺失：缺少关键文档扣5-10分
           - 功能缺陷：重要功能缺失每扣3-5分
        
        项目时间安排：
           - 需求分析：建议用时20%
           - 设计阶段：建议用时30%
           - 开发现：建议用35%
           - 测试优化：建议用时15%
        
        团队协作要求（如果是团队项目）：
           - 确定的任务分工
           - 期的进度汇报
           - 代码版本控制
           - 团队协作记录
        """,
        
        "期末试题": f"""
        {course_type}期末试题要求：
        
        试题难度要求（{config.get('difficulty', '中等')}）：
        {
            '''
            基础难度（Remember & Understand）：
            1. 题型分布：
               - 基础概念题（40%）：考察核心概念理解
               - 简单应用题（40%）：考察基本原理应用
               - 基础分析题（20%）：考察简单问题分析
            
            2. 知识覆盖：
               - 核心概念准确理解
               - 基本原理正确应用
               - 标准流程熟练掌握
               - 基础编程技能运用
            
            3. 答题要求：
               - 概念定义准确清晰
               - 基本步骤完整规范
               - 代码实现基本正确
               - 结果计算准确无误
            
            4. 评分标准：
               - 概念理解（40%）：术语准确、原理清晰
               - 应用能力（40%）：步骤正确、结果准确
               - 表达能力（20%）：条理清晰、书写规范
            ''' if config.get('difficulty') == '基础' else
            '''
            中等难度（Apply & Analyze）：
            1. 题型分布：
               - 概念应用题（30%）：考察知识综合运用
               - 综合分析题（40%）：考察问题分析能力
               - 实践应用题（30%）：考察实际问题解决
            
            2. 知识覆盖：
               - 多知识点综合运用
               - 实际问题分析能力
               - 算法设计与优化
               - 代码实现与调试
            
            3. 答题要求：
               - 分析过程逻辑完整
               - 解决方案合理可行
               - 代码实现规范高效
               - 结果分析深入准确
            
            4. 评分标准：
               - 分析能力（35%）：思路清晰、逻辑严谨
               - 解决方案（35%）：方案合理、实现正确
               - 创新思维（30%）：方法创新、优化改进
            ''' if config.get('difficulty') == '中等' else
            '''
            提高难度（Evaluate & Create）：
            1. 题型分布：
               - 综合应用题（30%）：考察系统性思维
               - 方案设计题（40%）：考察创新设计能力
               - 优化改进题（30%）：考察问题优化能力
            
            2. 知识覆盖：
               - 系统架构设计能力
               - 性能优化分析能力
               - 创新解决方案能力
               - 工程实践应用能力
            
            3. 答题要求：
               - 方案设计完整创新
               - 性能分析深入全面
               - 代码实现高效优雅
               - 优化思路清晰可行
            
            4. 评分标准：
               - 设计能力（30%）：方案创新、架构合理
               - 实现能力（30%）：代码优秀、性能优化
               - 创新能力（40%）：思路创新、解决方案优秀
            '''
        }
        
        考试基本要求：
        1. 考试时长：{config.get('duration', 120)}分钟
        2. 总分：{config.get('total_score', 100)}分
        
        题型要求（必须严格按照以下数量和分值生成）：
        {chr(10).join([
            f"   - {q_type}：必须生成{details['count']}题，每题{details['score']}分，总计{details['total']}分"
            for q_type, details in config.get('question_types', {}).items()
        ])}
        
        通用评分标准：
        1. 答题完整性（20%）：
           - 问题理解准确
           - 解答步骤完整
           - 结果表达清晰
        
        2. 专业规范性（30%）：
           - 专业术语使用准确
           - 解题思路专业规范
           - 代码编写符合规范
        
        3. 创新思维（20%）：
           - 解题思路创新
           - 方法选择合理
           - 优化改进意识
        
        4. 实践应用（30%）：
           - 理论联系实际
           - 解决实际问题
           - 考虑实际约束
        
        考试要求说明：
        1. 知识点覆盖：
           - 重点章节内容占比70%
           - 基础知识考察30%
           - 确保知识点分布合理
        
        2. 题目要求：
           - 每道题目必须标注对应的课程目标和AACSB目标
           - 每道题目必须包含详细的答案和解析
           - 题目描述必须清晰准确
           - 答案和评分标准必须明确
        
        3. 实施要求：
           - 题目数量必须严格按照要求生成
           - 每道题的分值必须与设定一致
           - 各种题型的总分必须符合设定
           - 必须覆盖选定的所有章节内容
           - 确保题目难度分布合理
           
        4. 特殊说明：
           - 选择题必须包含4个选项
           - 判断题必须明确是非
           - 简答题必须提供详细评分点
           - 编程题必须包含测试用例
        """
    }

    # 构建用户提示
    user_prompt = f"""
    课程基本信息：
    - 课程名称：{outline_data['basic_info']['course_name_cn']}
    - 课程代码：{outline_data['basic_info']['course_code']}
    - 课程类型：{course_type}
    - 开设院系：{department}
    - 专业：{major}
    - 考试类型：{exam_type}
    
    AACSB学习目标：
    {outline_data.get('aacsb_goals', '')}
    
    课程目标：
    {outline_data.get('course_objectives', '')}
    
    毕业要求：
    知识要求：
    {chr(10).join(outline_data['graduation_requirements']['knowledge'])}
    
    能力要求：
    {chr(10).join(outline_data['graduation_requirements']['ability'])}
    
    素质要求：
    {chr(10).join(outline_data['graduation_requirements']['quality'])}
    
    考核范围：
    """
    
    if chapters:
        user_prompt += "\n选定章节及关联：\n"
        for chapter in chapters:
            chapter_info = next((ch for ch in outline_data['course_schedule'] if ch['chapter'].startswith(str(chapter))), None)
            if chapter_info:
                # 获取前置和后续章节信息
                chapter_num = int(chapter_info['chapter'].split('.')[0])
                prev_chapter = next((ch for ch in outline_data['course_schedule'] if ch['chapter'].startswith(str(chapter_num-1))), None)
                next_chapter = next((ch for ch in outline_data['course_schedule'] if ch['chapter'].startswith(str(chapter_num+1))), None)
                
                user_prompt += f"""
                {chapter_info['chapter']}
                - 内容：{', '.join(chapter_info['content'])}
                - 要求：{', '.join(chapter_info['requirements'])}
                - 类型：{chapter_info['type']}
                - 知识关联：
                  * 前置知识：{prev_chapter['chapter'] + ': ' + ', '.join(prev_chapter['content']) if prev_chapter else '无'}
                  * 后续应用：{next_chapter['chapter'] + ': ' + ', '.join(next_chapter['content']) if next_chapter else '无'}
                  * 重点关联：{chapter_info.get('key_connections', '本章重点知识点的内在联系')}
                  * 实践应用：{chapter_info.get('practical_applications', '本章知识的实际应用场景')}
                """
    else:
        user_prompt += """
        覆盖全部章节：要求：
        1. 知识体系完整性：
           - 基础知识到高阶应用的递进
           - 理论知识到实践应用的转化
           - 不同章节知识点的有机结合
        
        2. 重点知识关联：
           - 核心概念的内在联系
           - 重点知识的应用场景
           - 难点知识的突破路径
        
        3. 能力培养路径：
           - 从基础理解到综合应用
           - 从简单实践到复杂项目
           - 从标准应用到创新设计
        """
    
    # 添加考试类型特定的提示
    exam_type_prompt = exam_type_prompts.get(exam_type, "")
    if exam_type_prompt:
        user_prompt += f"\n{exam_type}具体要求：\n{exam_type_prompt}"
    
    # 添加额外要求
    if additional_requirements:
        user_prompt += f"\n额外要求：\n{additional_requirements}"
    
    if exam_type == "实验":
        lab_type = config.get('lab_type', '综合性')
        lab_requirements = get_lab_type_requirements(lab_type, course_type, major)
        user_prompt += f"\n{lab_requirements}"
    
    return system_prompt, user_prompt

def generate_exam(outline_data, exam_type, chapters=None, additional_requirements=None, config=None, temperature=0.7):
    """调用DeepSeek API生成考试内容"""
    client = OpenAI(
        api_key=st.secrets["DEEPSEEK_API_KEY"],
        base_url="https://api.deepseek.com/beta"
    )
    
    system_prompt, user_prompt = create_exam_prompt(outline_data, exam_type, chapters, additional_requirements, config)
    
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=temperature  # 使用传入的temperature参数
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"API调用错误: {str(e)}")
        return None

def display_question(q, index):
    """显示题目内容"""
    # 定义难度显示的辅助函数
    def show_difficulty(difficulty):
        difficulty_colors = {
            "基础": "🟢",
            "中等": "🟡",
            "困难": "🔴",
            "基础性": "🟢",
            "综合性": "🟡",
            "设计性": "🔴"
        }
        return f"{difficulty_colors.get(difficulty, '🟡')} {difficulty}"

    if 'experiment' in q:  # 实验内容显示逻辑
        experiment = q['experiment']
        
        # 显示实验基本信息
        st.markdown(f"### {experiment.get('title', '实验')}")
        st.markdown(f"**实验类型：** {show_difficulty(experiment.get('type', '综合性实验'))}")
        st.markdown(f"**建议时长：** {experiment.get('duration', '4学时')}")
        
        # 显示实验目标
        if 'objectives' in experiment:
            st.markdown("### 实验目标")
            if 'knowledge' in experiment['objectives']:
                st.markdown("**知识目标：**")
                for goal in experiment['objectives']['knowledge']:
                    st.markdown(f"- {goal}")
            if 'skill' in experiment['objectives']:
                st.markdown("**技能目标：**")
                for goal in experiment['objectives']['skill']:
                    st.markdown(f"- {goal}")
            if 'course_objectives' in experiment['objectives']:
                st.markdown("**对应课程目标：**")
                for goal in experiment['objectives']['course_objectives']:
                    st.markdown(f"- {goal}")
            if 'aacsb_goals' in experiment['objectives']:
                st.markdown("**对应AACSB目标：**")
                for goal in experiment['objectives']['aacsb_goals']:
                    st.markdown(f"- {goal}")
        
        # 显示实验准备
        if 'prerequisites' in experiment:
            with st.expander("实验准备"):
                if 'knowledge' in experiment['prerequisites']:
                    st.markdown("**知识储备：**")
                    for k in experiment['prerequisites']['knowledge']:
                        st.markdown(f"- {k}")
                
                if 'environment' in experiment['prerequisites']:
                    st.markdown("**环境要求：**")
                    env_type_names = {
                        'hardware': '硬件要求',
                        'software': '软件要求',
                        'packages': '赖包'
                    }
                    for env_type, items in experiment['prerequisites']['environment'].items():
                        st.markdown(f"**{env_type_names.get(env_type, env_type)}：**")
                        for item in items:
                            st.markdown(f"- {item}")
                
                if 'references' in experiment['prerequisites']:
                    st.markdown("**参考资料：**")
                    for ref in experiment['prerequisites']['references']:
                        st.markdown(f"- {ref}")
        
        # 显示实验内容
        if 'content' in experiment:
            st.markdown("### 实验内容")
            if 'description' in experiment['content']:
                st.markdown(experiment['content']['description'])
            
            if 'steps' in experiment['content']:
                st.markdown("### 实验步骤")
                for step in experiment['content']['steps']:
                    with st.expander(f"步骤 {step['step_number']}: {step['title']}"):
                        st.markdown(step['description'])
                        if 'code_template' in step:
                            st.code(step['code_template'])
                        if 'expected_output' in step:
                            st.markdown("**预期输出：**")
                            st.markdown(step['expected_output'])
                        if 'notes' in step:
                            st.info(step['notes'])
        
        # 显示实要求
        if 'requirements' in experiment:
            with st.expander("实验要求"):
                requirement_names = {
                    'basic': '基本要求',
                    'good': '良好要求',
                    'excellent': '优秀要求',
                    'innovative': '创新要求'
                }
                for req_type, items in experiment['requirements'].items():
                    st.markdown(f"**{requirement_names.get(req_type, req_type)}：**")
                    for item in items:
                        st.markdown(f"- {item}")
        
        # 显示评分标准
        if 'grading_criteria' in experiment:
            with st.expander("评分标准"):
                criteria_names = {
                    'preparation': '实验准备',
                    'implementation': '实验实现',
                    'report': '实验报告'
                }
                for criterion, details in experiment['grading_criteria'].items():
                    st.markdown(f"**{criteria_names.get(criterion, criterion)}** ({details['weight']}分)")
                    for item in details['items']:
                        st.markdown(f"- {item['name']}({item['score']}分): {item['criteria']}")
        
        # 显示实验报告要求
        if 'report_template' in experiment:
            with st.expander("实验报告要求"):
                st.markdown("### 报告结构")
                if 'sections' in experiment['report_template']:
                    for section in experiment['report_template']['sections']:
                        st.markdown(f"### {section['title']}")
                        st.markdown(section['description'])
                        if 'requirements' in section:
                            for req in section['requirements']:
                                st.markdown(f"- {req}")
                        st.markdown("---")
                
                # 显示格式要求（如果有）
                if 'format_requirements' in experiment['report_template']:
                    st.markdown("### 格式要求")
                    format_reqs = experiment['report_template']['format_requirements']
                    
                    if 'general' in format_reqs:
                        st.markdown("**基本格式：**")
                        for req in format_reqs['general']:
                            st.markdown(f"- {req}")
                    
                    if 'content' in format_reqs:
                        st.markdown("**内容格式：**")
                        for req in format_reqs['content']:
                            st.markdown(f"- {req}")
                    
                    if 'submission' in format_reqs:
                        st.markdown("**提交要求：**")
                        for req in format_reqs['submission']:
                            st.markdown(f"- {req}")
        
    elif 'project' in q:  # 作业内容显示逻辑
        project = q['project']
        
        # 显示项目基本信息
        st.markdown(f"### {project.get('title', '大作业')}")
        if 'difficulty' in project:
            st.markdown(f"**难度级别：** {show_difficulty(project.get('difficulty', '中等'))}")
        st.markdown(f"**项目类型：** {project.get('type', '综合项目')}")
        st.markdown(f"**建议时长：** {project.get('duration', '4周')}")
        
        # 显示项目目标
        if 'objectives' in project:
            st.markdown("### 项目目标")
            for obj in project['objectives']:
                st.markdown(f"- {obj}")
        
        # 显示
        if 'requirements' in project:
            with st.expander("项目要求", expanded=True):
                for module, details in project['requirements'].items():
                    st.markdown(f"### {module}")
                    if '说明' in details:
                        st.markdown(f"**说明：** {details['说明']}")
                    if '交付物' in details:
                        st.markdown("**交付物：**")
                        for item in details['交付物']:
                            st.markdown(f"- {item}")
                    if '具体要求' in details:
                        st.markdown("**具体要求：**")
                        for item in details['具体要求']:
                            st.markdown(f"- {item}")
        
        # 显示评分标准
        if 'grading_criteria' in project:
            with st.expander("评分标准"):
                for criterion, details in project['grading_criteria'].items():
                    st.markdown(f"### {criterion}")
                    if '分值' in details:
                        st.markdown(f"**总分值：** {details['分值']}分")
                    if '评分项' in details:
                        for item in details['评分项']:
                            st.markdown(f"- {item['名称']}（{item['分数']}分）：{item['评分标准']}")
        
        # 显示提交要求
        if 'submission_requirements' in project:
            with st.expander("提交要求"):
                if isinstance(project['submission_requirements'], dict):
                    for req_type, requirements in project['submission_requirements'].items():
                        st.markdown(f"**{req_type}：**")
                        if isinstance(requirements, list):
                            for req in requirements:
                                st.markdown(f"- {req}")
                        else:
                            st.markdown(requirements)
                else:
                    for req in project['submission_requirements']:
                        st.markdown(f"- {req}")
        
        # 显示时间安排
        if 'timeline' in project:
            with st.expander("时间安排"):
                for phase, details in project['timeline'].items():
                    st.markdown(f"**{phase}：**")
                    if isinstance(details, list):
                        for item in details:
                            st.markdown(f"- {item}")
                    else:
                        st.markdown(details)
        
        # 示团队要求（如果有）
        if 'team_requirements' in project:
            with st.expander("团队要求"):
                for req in project['team_requirements']:
                    st.markdown(f"- {req}")
                    
    else:  # 普通题目显示逻辑
        st.markdown(f"### 第{index}题")
        
        # 使用列布局显示题目信息
        col1, col2, col3 = st.columns([1, 2, 2])
        with col1:
            if 'type' in q:
                st.markdown(f"**题型：** {q['type']}")
            if 'difficulty' in q:
                st.markdown(f"**难度：** {show_difficulty(q.get('difficulty', '中等'))}")
        
        with col2:
            if 'course_objectives' in q:
                objectives = q['course_objectives']
                if isinstance(objectives, list):
                    st.markdown(f"**课程目标：** {', '.join(objectives)}")
                else:
                    st.markdown(f"**课程目标：** {objectives}")
        
        with col3:
            if 'aacsb_goals' in q:
                goals = q['aacsb_goals']
                if isinstance(goals, list):
                    st.markdown(f"**AACSB目标：** {', '.join(goals)}")
                else:
                    st.markdown(f"**AACSB目：** {goals}")
        
        # 显示题目内容
        if 'question' in q:
            st.markdown("**题目：**")
            st.write(q['question'])
        
        # 显示选项（如果有）
        if 'options' in q and q['options']:
            st.markdown("**选项：**")
            for opt in q['options']:
                st.write(opt)
        
        # 显示答案和解析
        if 'answer' in q or 'explanation' in q:
            with st.expander("查看答案和解析"):
                if 'answer' in q:
                    st.markdown("**答案：**")
                    st.write(q['answer'])
                if 'explanation' in q:
                    st.markdown("**解析：**")
                    st.write(q['explanation'])

def get_project_json_format():
    """获取大作业的JSON格式"""
    return '''{
        "project": {
            "title": "项目标题",
            "type": "项目类型",
            "duration": "建议完成时间",
            "objectives": [
                "项目目标1",
                "项目目标2"
            ],
            "requirements": {
                "需求分析": {
                    "说明": "需求分析的详细描述",
                    "交付物": [
                        "需求规格说明书",
                        "用例文档"
                    ],
                    "具体要求": [
                        "具体要求1",
                        "具体要求2"
                    ]
                }
            },
            "grading_criteria": {
                "文档质量": {
                    "分值": 30,
                    "评分项": [
                        {"名称": "完整性", "分数": 10, "评分标准": "评分标准描述"},
                        {"名称": "规范性", "分数": 10, "评分标准": "评分标准描述"},
                        {"名称": "质量", "分数": 10, "评分标准": "评分标准描述"}
                    ]
                }
            },
            "submission_requirements": {
                "文档要求": [
                    "PDF格式的文档",
                    "Word格式的源文件"
                ],
                "代码要求": [
                    "源代码",
                    "可执行文件"
                ]
            },
            "timeline": {
                "第一周": [
                    "需求分析",
                    "概要设计"
                ],
                "第二周": [
                    "详细设计",
                    "始编码"
                ]
            },
            "team_requirements": [
                "团队规模：2-3人",
                "明确的分工"
            ]
        }
    }'''

def get_lab_type_requirements(lab_type, course_type, major):
    """获取不同类型实验的具体要求"""
    lab_types = {
        "基础": {
            "characteristics": [
                "针对单个知识点的掌握和应用",
                "提供详细的操作指导",
                "给出完整的代码框架",
                "设置明确的检查点",
                "预期结果清晰具体"
            ],
            "guidance_level": [
                "提供完整的实验步骤",
                "包含详细的代码示例",
                "给出具体的验证方法",
                "明确的完成标准"
            ],
            "evaluation_focus": [
                "基本概念的理解",
                "基本技能的掌握",
                "操作的规范性",
                "结果的正确"
            ],
            "report_requirements": [
                "详细记录实验步骤",
                "展示关键操作截图",
                "说明实验结果",
                "总结实验要点"
            ]
        },
        "综合性": {
            "characteristics": [
                "涉及多个知识点的综合应用",
                "需要综合运用多种技术",
                "包含一定的设计环节",
                "有一定的探索空间"
            ],
            "guidance_level": [
                "提供实",
                "给出关键步骤指导",
                "部分内容需要自主设计",
                "预期结果有一定弹性"
            ],
            "evaluation_focus": [
                "知识点的综合运用",
                "问题分析能力",
                "方案设计能力",
                "实现的完整性"
            ],
            "report_requirements": [
                "方案设计说明",
                "实现过程描述",
                "结果分析讨论",
                "创新点说明"
            ]
        },
        "设计性": {
            "characteristics": [
                "提供开放性问题",
                "需自主设计解决方案",
                "强调创新思维",
                "注重方案可行性"
            ],
            "guidance_level": [
                "只提供基本要",
                "生自主设计方案",
                "鼓励多样化解决方案",
                "重视创新性思维"
            ],
            "evaluation_focus": [
                "方案的创新性",
                "设计的合理性",
                "实现的可行性",
                "文档的专业性"
            ],
            "report_requirements": [
                "完整的设计文档",
                "详细的实现说明",
                "创新点分析",
                "改进方向建议"
            ]
        }
    }
    
    lab_type_info = lab_types.get(lab_type, lab_types["综合性"])
    
    return f"""
    {course_type}{lab_type}实验要求
    
    1. 实验特：
    {chr(10).join(f"   - {char}" for char in lab_type_info["characteristics"])}
    
    2. 指导方式：
    {chr(10).join(f"   - {guide}" for guide in lab_type_info["guidance_level"])}
    
    3. 评价重点：
    {chr(10).join(f"   - {eval}" for eval in lab_type_info["evaluation_focus"])}
    
    4. 报告求：
    {chr(10).join(f"   - {req}" for req in lab_type_info["report_requirements"])}
    
    5. 专业特色：
    - 结合{major}专业特点设计实验内容
    - 注重专业技能的培养
    - 符合行业实践要求
    """

def get_experiment_json_format():
    """获取实验内容的JSON格式"""
    return '''
    {
        "experiment": {
            "title": "实验标题",
            "type": "实验类型（基础性/综合性/设计性）",
            "duration": "建议实验时长",
            "objectives": {
                "knowledge": ["知识目标1", "知识目标2"],
                "skill": ["技能目标1", "技能目标2"],
                "course_objectives": ["对应的课程目标"],
                "aacsb_goals": ["对应的AACSB目标"]
            },
            "prerequisites": {
                "knowledge": ["前置知识要求"],
                "environment": {
                    "hardware": ["硬件要求"],
                    "software": ["软件要求"],
                    "packages": ["需要的包或库"]
                },
                "references": ["参考资料"]
            },
            "content": {
                "description": "实验内容概述",
                "steps": [
                    {
                        "step_number": "1",
                        "title": "步骤标题",
                        "description": "详细步骤说明",
                        "code_template": "相关代码模板",
                        "expected_output": "预期输出结果",
                        "notes": "注意事项"
                    }
                ]
            },
            "requirements": {
                "basic": ["基本要求（60分标准）"],
                "good": ["良好要求（80分标准）"],
                "excellent": ["优秀要求（90分标准）"],
                "innovative": ["创新要求（加分项）"]
            },
            "grading_criteria": {
                "preparation": {
                    "weight": 20,
                    "items": [
                        {"name": "环境配置", "score": 5, "criteria": "评分标准"},
                        {"name": "预习报告", "score": 15, "criteria": "评分标准"}
                    ]
                },
                "implementation": {
                    "weight": 50,
                    "items": [
                        {"name": "功能完成度", "score": 20, "criteria": "评分标准"},
                        {"name": "代码质量", "score": 15, "criteria": "评分标准"},
                        {"name": "实现效果", "score": 15, "criteria": "评分标准"}
                    ]
                },
                "report": {
                    "weight": 30,
                    "items": [
                        {"name": "实验报告", "score": 20, "criteria": "评分标准"},
                        {"name": "总结反思", "score": 10, "criteria": "评分标准"}
                    ]
                }
            },
            "report_template": {
                "sections": [
                    {
                        "title": "1. 实验目的",
                        "description": "明确说明本次实验要达到的目标，包括：",
                        "requirements": [
                            "理论知识点：需要掌握的核心概念和原理",
                            "实践技能点：需要培养的具体技能",
                            "创新目标：鼓励探索和创新的方向"
                        ]
                    },
                    {
                        "title": "2. 实验环境",
                        "description": "详细描述实验环境的配置情况，包括：",
                        "requirements": [
                            "硬件环境：计算机配置、特殊设备等",
                            "软件环境：操作系统、开发工具、版本信息等",
                            "相关依赖：需要安装的包、库、插件等",
                            "环境配置：具体的配置步骤和注意事项"
                        ]
                    },
                    {
                        "title": "3. 实验步骤",
                        "description": "按照时间顺序记录实验的完整过程，要求：",
                        "requirements": [
                            "前期准备：实验前的准备工作",
                            "操作过程：详细的操作步骤，配图说明",
                            "关键代码：核心代码片段及其说明",
                            "结果验证：每个步骤的验证方法和结果",
                            "注意事项：操作中的重点和难点"
                        ]
                    },
                    {
                        "title": "4. 实验结果",
                        "description": "展示和分析实验的最终结果，包括：",
                        "requirements": [
                            "结果展示：运行结果截图或输出",
                            "数据分析：相关数据的分析和说明",
                            "结果验证：验证结果的正确性",
                            "性能分析：程序性能相关数据（如有）",
                            "对比分析：与预期结果的对比"
                        ]
                    },
                    {
                        "title": "5. 问题与解决",
                        "description": "记录实验过程中遇到的问题及解决方案：",
                        "requirements": [
                            "问题描述：清晰描述遇到的问题",
                            "原因分析：分析问题产生的原因",
                            "解决过程：解决问题的具体步骤",
                            "解决方案：最终采用的解决方案",
                            "经验总结：解决问题的心得体会"
                        ]
                    },
                    {
                        "title": "6. 总结反思",
                        "description": "对本次实验进行总结和反思：",
                        "requirements": [
                            "知识总结：本次实验学到的主要知识点",
                            "技能提升：提升的具体技能",
                            "创新点：实验中的创新或改进之处",
                            "不足分析：实验中的不足之处",
                            "改进建议：对实验的改进建议",
                            "心得体会：个人感悟和建议"
                        ]
                    }
                ],
                "format_requirements": {
                    "general": [
                        "文档格式：PDF或Word格式",
                        "字体要求：正文宋体四，标题黑体三号",
                        "页边距：上下2.54cm，左右3.17cm",
                        "行间距：1.5倍行距",
                        "页码：页面底部居中"
                    ],
                    "content": [
                        "文字描述：清晰、准确、专业",
                        "图片要求：清晰、大小适中、有序号和说明",
                        "代码格式：统一的代码风格，有注释",
                        "引用标注：引用需注明来源",
                        "专业术语：准确使用专业术语"
                    ],
                    "submission": [
                        "文件命名：学号_姓名_实验X",
                        "提交方式：通过指定平台提交",
                        "提交时间：在规定截止日期前提交",
                        "相关材料：源代码等相关文件一并提交"
                    ]
                }
            }
        }
    }
    '''

def main():
    st.title("课程考试生成器 📚")
    st.markdown("---")
    
    # 上传课程大纲文件
    uploaded_file = st.file_uploader("上传课程大纲JSON文件", type=['json'])
    
    if uploaded_file is not None:
        outline_data = load_course_outline(uploaded_file)
        
        if outline_data:
            st.success("✅ 课程大纲加载成功！")
            # 保存课程名称到session state
            st.session_state.course_name = outline_data['basic_info']['course_name_cn']
            
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("课程信息")
                st.write(f"📘 课程名称：{outline_data['basic_info']['course_name_cn']}")
                st.write(f"🔢 课程代码：{outline_data['basic_info']['course_code']}")
                st.write(f"📝 考核方式：{outline_data['basic_info']['exam_type']}")
                
                # 显示课程简介
                with st.expander("查看课程简介"):
                    intro_keys = {
                        'position': '课程定位',
                        'purpose': '课程目的',
                        'content': '课程内容',
                        'method': '教学方法',
                        'outcome': '预期成果'
                    }
                    for key, value in outline_data['course_intro'].items():
                        st.markdown(f"**{intro_keys.get(key, key)}：** {value}")
                
                # 显示毕业要求
                with st.expander("查看毕业要求"):
                    st.markdown("**知识要求：**")
                    for item in outline_data['graduation_requirements']['knowledge']:
                        st.markdown(f"- {item}")
                    
                    st.markdown("**能力要求：**")
                    for item in outline_data['graduation_requirements']['ability']:
                        st.markdown(f"- {item}")
                    
                    st.markdown("**素质要求：**")
                    for item in outline_data['graduation_requirements']['quality']:
                        st.markdown(f"- {item}")
                
                # 显示AACSB目标
                with st.expander("查看AACSB目标"):
                    for goal in outline_data['aacsb_goals'].split('\n'):
                        if goal.strip():  # 确保不是空行
                            st.markdown(f"- {goal}")
                
                # 显示课程目标
                with st.expander("查看课程目标"):
                    for objective in outline_data['course_objectives'].split('\n'):
                        if objective.strip():  # 确保不是空行
                            st.markdown(f"- {objective}")
            
            with col2:
                # 将难度选择移到考试类型选择之前
                difficulty_level = st.radio(
                    "选择难度级别",
                    ["基础", "中等", "提高"],
                    help="""
                    难度级别说明：
                    
                    基础（Remember & Understand）：
                    • 基本概念理解和应用
                    • 单一知识点考察
                    • 步骤清晰，答案明确
                    • 适合：新手学习和基础巩固
                    
                    中等（Apply & Analyze）：
                    • 多知识点综合运用
                    • 需要一定分析和设计
                    • 包含实践应用场景
                    • 适合：有一定基础的学习者
                    
                    提高（Evaluate & Create）：
                    • 复杂问题解决
                    • 系统设计和优化
                    • 创新思维和实践
                    • 适合：进阶学习和能力提升
                    """,
                    horizontal=True
                )
                
                # 考试类型选择
                exam_types = ["练习", "实验", "大作业", "期末试题"]
                selected_type = st.selectbox(
                    "选择考核类型",
                    exam_types,
                    help="""
                    - 练习：用于日常练习和巩固知识点
                    - 实验：用于实验课程的实践内容
                    - 大作业：用于课程项目和综合实践
                    - 期末试题：用于期末考核
                    """
                )
                
                # 章节选择 - 为多选
                chapters = [
                    f"{ch['chapter']}"  # 直接使用chapter字段
                    for ch in outline_data['course_schedule']
                ]
                selected_chapters = st.multiselect(
                    "选择考核章", 
                    chapters,
                    default=[],  # 默认不选择任何章节
                    help="以选择个章节，不选择则默认覆所有章节"
                )

                # 根据选择的考试类型显示相关置
                if selected_type == "练习":
                    # 移除原有的难度选择
                    # 直接使用上面选择的 difficulty_level
                    
                    # 题型选择和配置
                    st.markdown("### 题型设置")
                    practice_types = {
                        "选择题": "考察基本概念理解",
                        "判断题": "考察知识点掌握",
                        "填空题": "考察关键知识点",
                        "简答题": "考察综合理解能力",
                        "编程题": "考察实践应用能力"
                    }
                    
                    # 选择题型
                    selected_practice_types = st.multiselect(
                        "选择题目类型",
                        list(practice_types.keys()),
                        help="""
                        可选题型说明：
                        • 选择题：适合考察基本概念理解和知识点掌握
                        • 判断题：适合考察关键知识点的准确理解
                        • 填空题：适合考察重要概念和关键术语
                        • 简答题：适合考察知识点的理解和应用
                        • 编程题：适合考察实践能力和综合应用
                        """
                    )
                    
                    # 如果有选择题型，显示配置选项
                    if selected_practice_types:
                        st.markdown("### 题型配置")
                        practice_config = {}
                        
                        for q_type in selected_practice_types:
                            st.markdown(f"**{q_type}**（{practice_types[q_type]}）")
                            count = st.number_input(
                                f"{q_type}数量",
                                min_value=1,
                                max_value=10,
                                value=2,
                                key=f"practice_count_{q_type}"
                            )
                            
                            practice_config[q_type] = {
                                "count": count,
                                "description": practice_types[q_type]
                            }
                        
                        # 显示题型统计
                        if practice_config:
                            st.markdown("### 练习统计")
                            total_questions = sum(config["count"] for config in practice_config.values())
                            
                            st.markdown(f"**总题数：** {total_questions}题")
                            
                            # 显示各题型统计
                            st.markdown("**各题型统计：**")
                            for q_type, config in practice_config.items():
                                st.markdown(f"- {q_type}：{config['count']}题（{config['description']}）")
                elif selected_type == "实验":
                    lab_type = st.selectbox(
                        "实验类型",
                        ["基础性", "综合性", "设计性"],
                        help="""
                        三种实验类型的特点：
                        
                        基础性实验：
                        • 针对单个知识点的掌握和应用
                        • 提供详细的实验指导和完整代码框架
                        • 设明确的检点和预期结果
                        • 适合：初学者，新概念学习
                        
                        综合性实验：
                        • 涉及多个知识点的综合应用
                        • 供关键步骤指导和部分代码框架
                        • 需要自主设计部分内容
                        • 适合：有一定基础的学习者
                        
                        设计性实验：
                        • 提供开放性问题
                        • 需要自主设计完整解决方案
                        • 强调创新思维和方案可行性
                        • 适合：基础扎实的高年级学生
                        """
                    )
                    
                    # 示所选实验类型的详细说明
                    lab_type_descriptions = {
                        "基础性": {
                            "目标": "掌握单个知识点和基本技能",
                            "特点": [
                                "• 详细的实验指导",
                                "• 完整的代码框架",
                                "• 明确的预期结果",
                                "• 规范的操作流程"
                            ],
                            "适用场景": [
                                "• 新概念的初次实践",
                                "• 基本技能的训练",
                                "• 标准流程的掌握"
                            ],
                            "评分标准": [
                                "• 操作的规范性（30%）",
                                "• 实验结果正确性（40%）",
                                "• 实验报告完整性（30%）"
                            ]
                        },
                        "综合性": {
                            "目标": "综合运用多个知识点解决问题",
                            "特点": [
                                "• 部分实验指导",
                                "• 框架代码+自主实现",
                                "• 一定的探索空间",
                                "• 需要方案设计"
                            ],
                            "适用场景": [
                                "• 多知识点的综合应用",
                                "• 中等难度问题解决",
                                "• 模块化功能实现"
                            ],
                            "评分标准": [
                                "• 方案设计合理性（30%）",
                                "• 功能实现完整性（40%）",
                                "• 实验报告质量（30%）"
                            ]
                        },
                        "设计性": {
                            "目标": "培养创新设计和工程实践能力",
                            "特点": [
                                "• 最小指导原则",
                                "• 完全自主设计",
                                "• 开放性问题",
                                "• 鼓励创新"
                            ],
                            "适用场景": [
                                "• 综合性项目实践",
                                "• 创新方案设计",
                                "• 实际问题解决"
                            ],
                            "评分标准": [
                                "• 设计创新性（30%）",
                                "• 实现可行性（40%）",
                                "• 文档专业性（30%）"
                            ]
                        }
                    }
                    
                    # 显示所选实验类型的详细说明
                    selected_lab_info = lab_type_descriptions[lab_type]
                    with st.expander(f"查看{lab_type}实验详细说明", expanded=True):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown(f"**实验目标：**\n{selected_lab_info['目标']}")
                            st.markdown("**实验特点：**")
                            for point in selected_lab_info['特点']:
                                st.markdown(point)
                        with col2:
                            st.markdown("**适用场景：**")
                            for scene in selected_lab_info['适用场景']:
                                st.markdown(scene)
                            st.markdown("**分标准：**")
                            for criterion in selected_lab_info['评分标准']:
                                st.markdown(criterion)
                elif selected_type == "大作业":
                    project_requirements = st.multiselect(
                        "项目要求",
                        ["需求分析", "概要设计", "详细设计", "代码实现", "测试报告"],
                        help="""
                        选择项目需要包含的内容：
                        - 需求分析：包含需求调研、用户需求分析、功能需求说明等
                        - 概要设计：包含系统架构、技术选型、模块划分等
                        - 详细设计：包含类图、ER图、数据库设计、API设计等
                        - 代码实现：包含源代码、注释文档、版本控制等
                        - 测试报告：包含测试计划、用例设计、执行记录等
                        
                        建议：
                        1. 完整项目建议选择3-4个模块
                        2. 小型项目可选择2-3个核心模块
                        3. 根据程重点和时间安排适当调整
                        """
                    )
                    
                    # 如果没有选择任何要求，显示提醒
                    if not project_requirements:
                        st.info("💡 请至少选择一个项目要求，建议选择多个模块以确保项目的完整性。", icon="ℹ️")
                elif selected_type == "期末试题":
                    # 考试基本设置
                    col21, col22 = st.columns(2)
                    with col21:
                        exam_duration = st.number_input("考试时长(分钟)", min_value=60, max_value=180, value=120, step=30)
                    with col22:
                        total_score = st.number_input("总分", min_value=60, max_value=100, value=100, step=10)
                    
                    # 题型选择和配置
                    st.markdown("### 题型设置")
                    question_types = {
                        "选择题": {"selected": False, "count": 0, "score": 0},
                        "判断题": {"selected": False, "count": 0, "score": 0},
                        "填空题": {"selected": False, "count": 0, "score": 0},
                        "简答题": {"selected": False, "count": 0, "score": 0},
                        "编程题": {"selected": False, "count": 0, "score": 0}
                    }
                    
                    # 选择题型
                    selected_types = st.multiselect(
                        "选择题目类型",
                        list(question_types.keys()),
                        help="择需要包型"
                    )
                    
                    # 如果有选择题型，显示配置选项
                    if selected_types:
                        st.markdown("### 题型配置")
                        current_total = 0
                        
                        for q_type in selected_types:
                            col1, col2 = st.columns(2)
                            with col1:
                                count = st.number_input(
                                    f"{q_type}数量",
                                    min_value=1,
                                    max_value=20,
                                    value=1,
                                    key=f"count_{q_type}"
                                )
                            with col2:
                                score = st.number_input(
                                    f"{q_type}每题分数",
                                    min_value=1,
                                    max_value=50,
                                    value=5,
                                    key=f"score_{q_type}"
                                )
                            question_types[q_type]["selected"] = True
                            question_types[q_type]["count"] = count
                            question_types[q_type]["score"] = score
                            current_total += count * score
                        
                        # 显示总分统计
                        st.markdown("### 分数统计")
                        st.markdown(f"当前总分：{current_total}")
                        
                        # 检查总分是否符合要求
                        if current_total != total_score:
                            st.warning(f"⚠️ 当前配置的总分（{current_total}分）与设定的总分（{total_score}分）不符！请调整题目数量或分值。")
                            
                            # 显示每种题型的总分
                            st.markdown("### 各题型得分：")
                            for q_type in selected_types:
                                type_total = question_types[q_type]["count"] * question_types[q_type]["score"]
                                st.markdown(f"- {q_type}：{question_types[q_type]['count']}题 × {question_types[q_type]['score']}分 = {type_total}分")

            # 添加额外要求输入框
            st.markdown("### 额外生成要求")
            
            # 根据不同考试类型提供不同的placeholder提示
            placeholders = {
                "练习": "例如\n1. 针对特定知识点\n2. 包含解题技巧\n3. 难度要求",
                "实验": "例如：\n1. 实验环境要求\n2. 具体实验步骤\n3. 验收标准",
                "大作业": "例如：\n1. 项目具体要求\n2. 技术栈限制\n3. 创新点要求",
                "期末试题": "\n1. 题型分布\n2. 难度分布\n3. 重考察内容"
            }
            
            additional_requirements = st.text_area(
                "请输入额外的生成要求（可选）",
                placeholder=placeholders.get(selected_type, "请输入额外要求"),
                help="在这里输入任何额外的要求，这些要求将被用于定制生成的考试内容"
            )
            
            # 在 main 函数中修改生成内容的部分
            if st.button("🎯 生成考试内容", use_container_width=True):
                # 构建配置信息
                config = {
                    "type": selected_type,
                    "difficulty": difficulty_level,  # 添加通用难度设置
                    "chapters": selected_chapters if selected_chapters else None
                }
                
                # 根据不同考试类型添加特定配置
                if selected_type == "练习":
                    if selected_practice_types:
                        config["practice_types"] = practice_config
                elif selected_type == "实验":
                    config["lab_type"] = lab_type
                elif selected_type == "大作业":
                    config["project_requirements"] = project_requirements
                elif selected_type == "期末试题":
                    config["duration"] = exam_duration
                    config["total_score"] = total_score
                    config["question_types"] = {
                        q_type: {
                            "count": question_types[q_type]["count"],
                            "score": question_types[q_type]["score"],
                            "total": question_types[q_type]["count"] * question_types[q_type]["score"]
                        }
                        for q_type in selected_types
                    }
                
                # 保存生成的配置到session state
                st.session_state.last_config = config
                
                # 增加一个temperature参数到session state
                if 'temperature' not in st.session_state:
                    st.session_state.temperature = 0.7
                
                with st.spinner("正在生成考试内容，请稍候..."):
                    exam_content = generate_exam(
                        outline_data, 
                        selected_type,  # 使用 selected_type 而不是 config["type"]
                        selected_chapters,  # 使用 selected_chapters 而不是 config["chapters"]
                        additional_requirements,
                        config,
                        temperature=st.session_state.temperature
                    )
                    
                    # 保存生成的内容到session state
                    st.session_state.last_exam_content = exam_content
                    
                    # 显示生成的内容
                    display_exam_content(exam_content, selected_type)

            # 添加重新生成按钮
            if 'last_config' in st.session_state:
                if st.button("🔄 重新生成", use_container_width=True):
                    with st.spinner("正在重新生成内容，请稍候..."):
                        # 增加temperature以增加随机性
                        st.session_state.temperature += 0.1
                        if st.session_state.temperature > 1.0:
                            st.session_state.temperature = 0.7
                        
                        # 构建新的提示词
                        additional_reqs = f"请生成与之前不同的内容。当前随机性参数：{st.session_state.temperature}"
                        
                        # 重新生成内容
                        new_exam_content = generate_exam(
                            outline_data,
                            st.session_state.last_config["type"],
                            st.session_state.last_config.get("chapters"),  # 使用 get 方法避免 KeyError
                            additional_reqs,
                            st.session_state.last_config,
                            temperature=st.session_state.temperature
                        )
                        
                        # 保存新生成的内容到session state
                        st.session_state.last_exam_content = new_exam_content
                        
                        # 显示新生成的内容
                        display_exam_content(new_exam_content, selected_type)

            # 添加下载按钮部分
            if 'last_exam_content' in st.session_state:
                st.markdown("### 下载选项")
                col1, col2 = st.columns(2)
                
                with col1:
                    # JSON格式下载
                    json_data = st.session_state.last_exam_content
                    if st.download_button(
                        label="📥 下载JSON格式",
                        data=json.dumps(json_data, ensure_ascii=False, indent=2),
                        file_name=f"{st.session_state.course_name}_{selected_type}.json",
                        mime="application/json",
                        help="下载JSON格式的原始数据",
                        use_container_width=True
                    ):
                        # 同时保存到服务器
                        success, result = save_json_to_server(
                            json_data,
                            outline_data['basic_info']['course_name_cn'],
                            outline_data['basic_info']['department'],
                            outline_data['basic_info']['major'],
                            selected_type
                        )
                        if success:
                            st.success(f"文件已保存到服务器: {result}")
                        else:
                            st.error(f"保存到服务器失败: {result}")
                
                with col2:
                    # Word格式下载
                    doc_io = create_word_document(st.session_state.last_exam_content, selected_type, st.session_state.course_name)
                    st.download_button(
                        label="📄 下载Word格式",
                        data=doc_io.getvalue(),
                        file_name=f"{st.session_state.course_name}_{selected_type}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="下载Word格式的文档",
                        use_container_width=True
                    )

def create_word_document(exam_content, selected_type, course_name):
    """创建Word文档"""
    doc = Document()
    
    # 创建代码样式
    styles = doc.styles
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(10)
        # 添加其他代码样式的格式设置
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
    except ValueError:
        # 如果样式已存在，就使用现有的样式
        code_style = styles['Code']

    # 设置标题
    title = doc.add_heading(f'{course_name} - {selected_type}', 0)
    
    if selected_type == "大作业":
        if 'project' in exam_content:
            project = exam_content['project']
            
            # 添加项目标题
            doc.add_heading(project.get('title', '大作业'), level=1)
            
            # 添加项目基本信息
            doc.add_paragraph(f"项目类型：{project.get('type', '综合项目')}")
            doc.add_paragraph(f"建议��成时间：{project.get('duration', '4周')}")
            
            # 添加项目目标
            if 'objectives' in project:
                doc.add_heading('项目目标', level=2)
                for obj in project['objectives']:
                    doc.add_paragraph(obj, style='List Bullet')
            
            # 添加项目要求
            if 'requirements' in project:
                doc.add_heading('项目要求', level=2)
                for module, details in project['requirements'].items():
                    doc.add_heading(module, level=3)
                    if '说明' in details:
                        doc.add_paragraph(f"说明：{details['说明']}")
                    if '交付物' in details:
                        doc.add_paragraph("交付物：")
                        for item in details['交付物']:
                            doc.add_paragraph(item, style='List Bullet')
                    if '具体要求' in details:
                        doc.add_paragraph("具体要求：")
                        for item in details['具体要求']:
                            doc.add_paragraph(item, style='List Bullet')
            
            # 添加评分标准
            if 'grading_criteria' in project:
                doc.add_heading('评分标准', level=2)
                for criterion, details in project['grading_criteria'].items():
                    doc.add_heading(criterion, level=3)
                    if '分值' in details:
                        doc.add_paragraph(f"总分值：{details['分值']}分")
                    if '评分项' in details:
                        for item in details['评分项']:
                            doc.add_paragraph(f"- {item['名称']}（{item['分数']}分）：{item['评分标准']}")
            
            # 添加提交要求
            if 'submission_requirements' in project:
                doc.add_heading('提交要求', level=2)
                if isinstance(project['submission_requirements'], dict):
                    for req_type, requirements in project['submission_requirements'].items():
                        doc.add_paragraph(f"{req_type}：")
                        if isinstance(requirements, list):
                            for req in requirements:
                                doc.add_paragraph(req, style='List Bullet')
                        else:
                            doc.add_paragraph(requirements)
                else:
                    for req in project['submission_requirements']:
                        doc.add_paragraph(req, style='List Bullet')
            
            # 添加时间安排
            if 'timeline' in project:
                doc.add_heading('时间安排', level=2)
                for phase, details in project['timeline'].items():
                    doc.add_paragraph(f"{phase}：")
                    if isinstance(details, list):
                        for item in details:
                            doc.add_paragraph(item, style='List Bullet')
                    else:
                        doc.add_paragraph(details)
            
            # 添加团队要求
            if 'team_requirements' in project:
                doc.add_heading('团队要求', level=2)
                for req in project['team_requirements']:
                    doc.add_paragraph(req, style='List Bullet')
    
    elif selected_type == "实验":
        if 'experiment' in exam_content:
            experiment = exam_content['experiment']
            
            # 添加实验标题
            doc.add_heading(experiment.get('title', '实验'), level=1)
            
            # 添加实验类型和时长
            doc.add_paragraph(f"实验类型：{experiment.get('type', '综合性实验')}")
            doc.add_paragraph(f"建议时长：{experiment.get('duration', '4学时')}")
            
            # 添加实验目标
            if 'objectives' in experiment:
                doc.add_heading('实验目标', level=2)
                if 'knowledge' in experiment['objectives']:
                    doc.add_paragraph('知识目标：')
                    for goal in experiment['objectives']['knowledge']:
                        doc.add_paragraph(goal, style='List Bullet')
                if 'skill' in experiment['objectives']:
                    doc.add_paragraph('技能目标：')
                    for goal in experiment['objectives']['skill']:
                        doc.add_paragraph(goal, style='List Bullet')
                if 'course_objectives' in experiment['objectives']:
                    doc.add_paragraph('对应课程目标：')
                    for goal in experiment['objectives']['course_objectives']:
                        doc.add_paragraph(goal, style='List Bullet')
                if 'aacsb_goals' in experiment['objectives']:
                    doc.add_paragraph('对应AACSB目标：')
                    for goal in experiment['objectives']['aacsb_goals']:
                        doc.add_paragraph(goal, style='List Bullet')
            
            # 添加实验准备
            if 'prerequisites' in experiment:
                doc.add_heading('实验准备', level=2)
                if 'knowledge' in experiment['prerequisites']:
                    doc.add_paragraph('知识储备：')
                    for k in experiment['prerequisites']['knowledge']:
                        doc.add_paragraph(k, style='List Bullet')
                
                if 'environment' in experiment['prerequisites']:
                    doc.add_paragraph('环境要求：')
                    env_type_names = {
                        'hardware': '硬件要求',
                        'software': '软件要求',
                        'packages': '依赖包'
                    }
                    for env_type, items in experiment['prerequisites']['environment'].items():
                        doc.add_paragraph(f"{env_type_names.get(env_type, env_type)}：")
                        for item in items:
                            doc.add_paragraph(item, style='List Bullet')
                
                if 'references' in experiment['prerequisites']:
                    doc.add_paragraph('参考资料：')
                    for ref in experiment['prerequisites']['references']:
                        doc.add_paragraph(ref, style='List Bullet')
            
            # 添加实验内容
            if 'content' in experiment:
                doc.add_heading('实验内容', level=2)
                if 'description' in experiment['content']:
                    doc.add_paragraph(experiment['content']['description'])
                
                if 'steps' in experiment['content']:
                    doc.add_heading('实验步骤', level=2)
                    for step in experiment['content']['steps']:
                        doc.add_heading(f"步骤 {step['step_number']}: {step['title']}", level=3)
                        doc.add_paragraph(step['description'])
                        if 'code_template' in step:
                            doc.add_paragraph('代码模板：')
                            try:
                                # 尝试使用Code样式
                                doc.add_paragraph(step['code_template'], style='Code')
                            except KeyError:
                                # 如果样式应用失败，使用普通段落
                                p = doc.add_paragraph(step['code_template'])
                                p.paragraph_format.left_indent = Inches(0.5)
                                run = p.runs[0]
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                        if 'expected_output' in step:
                            doc.add_paragraph('预期输出：')
                            doc.add_paragraph(step['expected_output'])
                        if 'notes' in step:
                            doc.add_paragraph('注意事项：')
                            doc.add_paragraph(step['notes'])
            
            # 添加评分标准
            if 'grading_criteria' in experiment:
                doc.add_heading('评分标准', level=2)
                for criterion, details in experiment['grading_criteria'].items():
                    doc.add_paragraph(f"{criterion}（{details['weight']}分）：")
                    for item in details['items']:
                        doc.add_paragraph(f"- {item['name']}({item['score']}分): {item['criteria']}")
            
            # 添加实验报告要求
            if 'report_template' in experiment:
                doc.add_heading('实验报告要求', level=2)
                if 'sections' in experiment['report_template']:
                    for section in experiment['report_template']['sections']:
                        doc.add_heading(section['title'], level=3)
                        doc.add_paragraph(section['description'])
                        if 'requirements' in section:
                            for req in section['requirements']:
                                doc.add_paragraph(req, style='List Bullet')
                
                if 'format_requirements' in experiment['report_template']:
                    doc.add_heading('格式要求', level=3)
                    format_reqs = experiment['report_template']['format_requirements']
                    
                    if 'general' in format_reqs:
                        doc.add_paragraph('基本格式：')
                        for req in format_reqs['general']:
                            doc.add_paragraph(req, style='List Bullet')
                    
                    if 'content' in format_reqs:
                        doc.add_paragraph('内容格式：')
                        for req in format_reqs['content']:
                            doc.add_paragraph(req, style='List Bullet')
                    
                    if 'submission' in format_reqs:
                        doc.add_paragraph('提交要求：')
                        for req in format_reqs['submission']:
                            doc.add_paragraph(req, style='List Bullet')
    
    else:  # 练习或期末试题
        if 'questions' in exam_content:
            for i, q in enumerate(exam_content['questions'], 1):
                # 添加题目标题
                doc.add_heading(f'第{i}题 ({q["type"]})', level=2)
                
                # 添加题目内容
                doc.add_paragraph(q['question'])
                
                # 添加选项（如果有）
                if 'options' in q and q['options']:
                    for opt in q['options']:
                        doc.add_paragraph(opt)
                
                # 添加答案和解析
                if 'answer' in q or 'explanation' in q:
                    doc.add_heading('答案和解析：', level=3)
                    if 'answer' in q:
                        doc.add_paragraph(f'答案：{q["answer"]}')
                    if 'explanation' in q:
                        doc.add_paragraph(f'解析：{q["explanation"]}')
                
                doc.add_paragraph('') # 添加空行分隔
    
    # 保存到内存
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def display_exam_content(exam_content, selected_type):
    """显示生成的考试内容"""
    if exam_content:
        if selected_type == "大作业":
            st.markdown("---")
            st.subheader("📝 生成的大作业内容")
            if 'project' in exam_content:
                display_question(exam_content, 1)
            else:
                st.error("生成的大作业内容格式不正确")
        elif selected_type == "实验":
            st.markdown("---")
            st.subheader("📝 生成的实验内容")
            if 'experiment' in exam_content:
                display_question(exam_content, 1)
            else:
                st.error("生成的实验内容格式不正确")
        else:
            st.markdown("---")
            st.subheader("📝 生成的考试内容")
            if 'questions' in exam_content:
                for i, q in enumerate(exam_content['questions'], 1):
                    display_question(q, i)
                    st.markdown("---")
            else:
                st.error("生成的考试内容格式不正确")
    else:
        st.error("未能成功生成内容")

def save_json_to_server(json_data, course_name, department, major, exam_type):
    """保存JSON文件到服务器指定目录"""
    # 构建目录路径
    dir_path = f"{course_name} - {department} - {major}"
    
    # 创建目录（如果不存在）
    os.makedirs(dir_path, exist_ok=True)
    
    # 生成文件名（包含时间戳以避免重名）
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{course_name}_{exam_type}_{timestamp}.json"
    file_path = os.path.join(dir_path, file_name)
    
    # 保存文件
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        return True, file_path
    except Exception as e:
        return False, str(e)

if __name__ == "__main__":
    main()
